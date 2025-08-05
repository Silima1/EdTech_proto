import streamlit as st
import pandas as pd
import tempfile
import os
import docx
from PyPDF2 import PdfReader
from fpdf import FPDF
import spacy
from dotenv import load_dotenv
from langchain.chat_models import ChatOpenAI
import json

load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

llm = ChatOpenAI(model="gpt-4", temperature=0.2, openai_api_key=openai_api_key)
nlp = spacy.load("pt_core_news_sm")


def extract_name_from_first_page(content):
    doc = nlp(content)
    for ent in doc.ents:
        if ent.label_ == "PER":
            return ent.text
    return "Nome Não Identificado"


def load_docx_first_page(file_path):
    doc = docx.Document(file_path)
    return doc.paragraphs[0].text if doc.paragraphs else "Nenhum conteúdo encontrado"


def process_document(file):
    temp_filepath = f"./temp_{file.name}"
    with open(temp_filepath, "wb") as temp_file:
        temp_file.write(file.getvalue())
    if file.type == "application/pdf":
        reader = PdfReader(temp_filepath)
        content = (
            reader.pages[0].extract_text()
            if len(reader.pages) > 0
            else "Nenhum conteúdo encontrado"
        )
        documents = [{"content": content, "source": file.name}]
    elif file.type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]:
        content = load_docx_first_page(temp_filepath)
        documents = [{"content": content, "source": file.name}]
    else:
        documents = None
    os.remove(temp_filepath)
    return documents


def evaluate_sources(documents):
    results = []
    for doc in documents:
        content = doc["content"]
        name = extract_name_from_first_page(content)

        prompt = f"""
Você é um agente avaliador de trabalhos universitários da área de Engenharia Informática.

Avalie o conteúdo abaixo e responda **apenas** com um objeto JSON no seguinte formato:

{{
  "nota": número entre 0 e 20,
  "feedback": "texto explicativo e objetivo como assistente automático"
}}

Conteúdo:
\"\"\"
{content}
\"\"\"
"""
        try:
            response = llm.predict(prompt)
            parsed = json.loads(response)

            grade = int(parsed.get("nota", 10))
            feedback = parsed.get("feedback", "").strip()

            results.append({
                "name": name,
                "grade": grade,
                "feedback": feedback,
                "plagiarism": "IA textual",
            })
        except Exception as e:
            results.append({
                "name": name,
                "grade": 0,
                "feedback": f"Erro ao processar com IA: {e}",
                "plagiarism": "Erro",
            })
    return results


def generate_csv_report(data, curso, ano, nivel):
    df = pd.DataFrame(data)
    csv_file = tempfile.NamedTemporaryFile(delete=False, suffix=".csv")
    df.to_csv(csv_file.name, index=False, encoding="utf-8-sig")
    return csv_file.name


def generate_pdf_report(data, curso, ano, nivel):
    pdf_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.cell(200, 10, txt="Relatório de Correção", ln=True, align="C")
    pdf.cell(200, 10, txt=f"Curso: {curso} | Nível: {nivel} | Ano: {ano}", ln=True, align="C")
    pdf.ln(10)

    for entry in data:
        pdf.cell(200, 10, txt=f"Nome: {entry['Nome do Estudante']}", ln=True)
        pdf.cell(200, 10, txt=f"Nota: {entry['Nota']}", ln=True)
        pdf.cell(200, 10, txt=f"Feedback: {entry['Feedback']}", ln=True)
        pdf.cell(200, 10, txt=f"Percentual de Plágio: {entry['Percentual de Plágio']}", ln=True)
        pdf.ln(5)

    pdf.output(pdf_file.name)
    return pdf_file.name


# Streamlit UI
st.set_page_config(page_title="Professor Virtual 📚", page_icon="📚", layout="wide")
st.title("Professor Virtual 📚")
st.sidebar.header("Configurações")

curso = st.sidebar.selectbox("Selecione o Curso", ["Engenharia Informática", "Direito", "Medicina", "Arquitetura"])
ano = st.sidebar.selectbox("Selecione o Ano Acadêmico", ["1º Ano", "2º Ano", "3º Ano", "Mestrado"])
nivel = st.sidebar.selectbox("Selecione o Nível", ["Licenciatura", "Mestrado", "Doutorado"])

st.sidebar.header("Envio de Trabalhos")
uploads = st.sidebar.file_uploader("Envie os documentos (PDF ou DOCX)", type=["pdf", "docx"], accept_multiple_files=True)

if not uploads:
    st.info("Por favor, envie documentos para começar.")
    st.stop()

documents = []
for upload in uploads:
    processed_docs = process_document(upload)
    if processed_docs:
        documents.extend(processed_docs)

if not documents:
    st.error("Nenhum documento válido encontrado.")
    st.stop()

results = evaluate_sources(documents)

report_data = []
for result in results:
    report_data.append(
        {
            "Nome do Estudante": result["name"],
            "Curso": curso,
            "Nível": nivel,
            "Ano": ano,
            "Nota": result["grade"],
            "Feedback": result["feedback"],
            "Percentual de Plágio": f"{result['plagiarism']}",
        }
    )

csv_file = generate_csv_report(report_data, curso, ano, nivel)
pdf_file = generate_pdf_report(report_data, curso, ano, nivel)

st.write("### Resultados da Avaliação")
for entry in report_data:
    st.write(f"**Nome:** {entry['Nome do Estudante']}")
    st.write(f"**Curso:** {entry['Curso']}")
    st.write(f"**Nível:** {entry['Nível']}")
    st.write(f"**Ano:** {entry['Ano']}")
    st.write(f"**Nota:** {entry['Nota']}")
    st.write(f"**Feedback:** {entry['Feedback']}")
    st.write(f"**Percentual de Plágio:** {entry['Percentual de Plágio']}")
    st.write("---")

st.download_button("Pauta", data=open(csv_file, "rb"), file_name="pauta.csv", mime="text/csv")
st.download_button("Relatório de Correção", data=open(pdf_file, "rb"), file_name="relatorio.pdf", mime="application/pdf")

if st.button("Explorar Recursos do Moodle"):
    st.write("Explorando os recursos do Moodle para correção de quizzes...")

try:
    os.remove(csv_file)
    os.remove(pdf_file)
except Exception as e:
    st.warning(f"Erro ao limpar arquivos temporários: {e}")